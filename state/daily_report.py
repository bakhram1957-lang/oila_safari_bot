"""
Oilaviy safar rejasi - kunlik bot.

Bu skript GitHub Actions orqali har kuni (taxminan soat 20:00, Toshkent vaqti) ishga tushadi:
1. Telegram bot'ga kelgan yangi xabarlarni o'qiydi (getUpdates).
2. "Vaqt / Ism / Qayerdan / Qayerga" formatidagi xabarlarni ajratib oladi.
3. Bir xil vaqt+yo'nalishdagi safarlarni birlashtiradi.
4. Har bir safarni 2 mashinadan biriga (vaqt to'qnashuvini hisobga olib) taqsimlaydi.
5. Excel va Word fayllarini tayyorlaydi va asosiy qabul qiluvchiga (owner) Telegram orqali yuboradi.
6. Holatni (oxirgi update_id, owner_chat_id) state/config.json fayliga saqlaydi.

Muhit o'zgaruvchisi: TELEGRAM_BOT_TOKEN (GitHub repository secret sifatida saqlanadi).
"""

import os
import json
import re
from datetime import datetime, timedelta, timezone

import requests

STATE_PATH = os.path.join(os.path.dirname(__file__), "state", "config.json")
TASHKENT = timezone(timedelta(hours=5))

TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN")
API = f"https://api.telegram.org/bot{TOKEN}" if TOKEN else None

TIME_RE = re.compile(r"^([01]?\d|2[0-3]):([0-5]\d)$")


# --------------------------------------------------------------------------
# Holatni o'qish / yozish
# --------------------------------------------------------------------------

def load_state():
    with open(STATE_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def save_state(state):
    with open(STATE_PATH, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)


# --------------------------------------------------------------------------
# Telegram API yordamchilari
# --------------------------------------------------------------------------

def api_call(method, **params):
    r = requests.post(f"{API}/{method}", data=params, timeout=30)
    r.raise_for_status()
    return r.json()


def send_message(chat_id, text):
    try:
        api_call("sendMessage", chat_id=chat_id, text=text)
    except Exception as e:  # noqa: BLE001
        print(f"sendMessage xato (chat_id={chat_id}): {e}")


def send_document(chat_id, file_path, caption=None):
    with open(file_path, "rb") as f:
        files = {"document": f}
        data = {"chat_id": chat_id}
        if caption:
            data["caption"] = caption
        r = requests.post(f"{API}/sendDocument", data=data, files=files, timeout=60)
        r.raise_for_status()


def get_updates(offset):
    resp = api_call("getUpdates", offset=offset + 1, timeout=0)
    return resp.get("result", [])


# --------------------------------------------------------------------------
# Xabarlarni tahlil qilish
# --------------------------------------------------------------------------

def parse_trip_message(text):
    """4 qatorli xabarni (Vaqt / Ism / Qayerdan / Qayerga) ajratadi."""
    lines = [l.strip() for l in text.strip().splitlines() if l.strip()]
    if len(lines) != 4:
        return None
    vaqt, ism, qayerdan, qayerga = lines
    if not TIME_RE.match(vaqt):
        return None
    names = [n.strip() for n in ism.split(",") if n.strip()]
    if not names or not qayerdan or not qayerga:
        return None
    return {"vaqt": vaqt, "ismlar": names, "qayerdan": qayerdan, "qayerga": qayerga}


def merge_trips(trips):
    """Bir xil vaqt+yo'nalishdagi safarlarni bitta yozuvga birlashtiradi."""
    groups = {}
    order = []
    for t in trips:
        key = (t["vaqt"], t["qayerdan"], t["qayerga"])
        if key not in groups:
            groups[key] = {
                "vaqt": t["vaqt"],
                "qayerdan": t["qayerdan"],
                "qayerga": t["qayerga"],
                "ismlar": [],
            }
            order.append(key)
        for name in t["ismlar"]:
            if name not in groups[key]["ismlar"]:
                groups[key]["ismlar"].append(name)
    merged = [groups[k] for k in order]
    merged.sort(key=lambda x: x["vaqt"])
    return merged


# --------------------------------------------------------------------------
# Mashinalarga taqsimlash
# --------------------------------------------------------------------------

def time_to_minutes(t):
    h, m = t.split(":")
    return int(h) * 60 + int(m)


def assign_cars(trips, cars, duration_minutes):
    """
    Har bir safarni ikki mashinadan biriga taqsimlaydi.
    - Vaqt jihatdan bo'sh bo'lgan mashina tanlanadi (eng kam band bo'lgani ustunlik oladi).
    - Ikkala mashina ham band bo'lsa -> "ziddiyat" deb belgilanadi (qo'lda tekshirish kerak).
    - Yo'lovchilar soni mashina sig'imidan oshsa -> ogohlantirish qo'shiladi.
    """
    car_state = [{"free_from": 0, "load": 0} for _ in cars]
    results = []
    for trip in trips:
        start = time_to_minutes(trip["vaqt"])
        end = start + duration_minutes

        free_cars = [i for i, cs in enumerate(car_state) if cs["free_from"] <= start]
        notes = []

        if free_cars:
            chosen = min(free_cars, key=lambda i: car_state[i]["load"])
        else:
            chosen = min(range(len(car_state)), key=lambda i: car_state[i]["free_from"])
            notes.append("DIQQAT: vaqt to'qnashuvi, qo'lda tekshiring")

        passenger_count = len(trip["ismlar"])
        cap = cars[chosen].get("adult_capacity", 4) + cars[chosen].get("child_capacity", 0)
        if passenger_count > cap:
            notes.append("DIQQAT: yo'lovchi soni sig'imdan ko'p bo'lishi mumkin")

        car_state[chosen]["free_from"] = end
        car_state[chosen]["load"] += 1

        results.append({**trip, "mashina_idx": chosen, "izoh": "; ".join(notes)})
    return results


# --------------------------------------------------------------------------
# Fayllarni tayyorlash
# --------------------------------------------------------------------------

def build_excel(assigned, cars, sana, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Reja"

    headers = ["Vaqt", "Ism(lar)", "Qayerdan", "Qayerga", "Mashina", "Shofyor", "Izoh"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="2F5597", end_color="2F5597", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    for t in assigned:
        car = cars[t["mashina_idx"]]
        ws.append(
            [
                t["vaqt"],
                ", ".join(t["ismlar"]),
                t["qayerdan"],
                t["qayerga"],
                f"{car['name']} ({car['model']})",
                car["driver"],
                t["izoh"],
            ]
        )

    widths = [8, 24, 16, 16, 24, 22, 34]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + i)].width = w
    ws.freeze_panes = "A2"

    ws.title = f"Reja {sana}"
    wb.save(path)


def build_word(assigned, cars, sana, path):
    from docx import Document

    doc = Document()
    doc.add_heading(f"Oilaviy safar rejasi - {sana}", level=1)

    if not assigned:
        doc.add_paragraph("Ertangi kun uchun hech qanday so'rov kelmadi.")
        doc.save(path)
        return

    for idx, car in enumerate(cars):
        doc.add_heading(f"{car['name']}: {car['driver']} - {car['model']}", level=2)
        car_trips = [t for t in assigned if t["mashina_idx"] == idx]
        if not car_trips:
            doc.add_paragraph("Bu mashina uchun reja yo'q.")
            continue

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Vaqt", "Ism(lar)", "Qayerdan", "Qayerga", "Izoh"]):
            hdr[i].text = h

        for t in car_trips:
            row = table.add_row().cells
            row[0].text = t["vaqt"]
            row[1].text = ", ".join(t["ismlar"])
            row[2].text = t["qayerdan"]
            row[3].text = t["qayerga"]
            row[4].text = t["izoh"]

    doc.save(path)


# --------------------------------------------------------------------------
# Asosiy oqim
# --------------------------------------------------------------------------

def process_updates(state, updates):
    """Yangi xabarlarni qayta ishlaydi, state'ni va trips ro'yxatini qaytaradi."""
    trips = []
    max_update_id = state.get("last_update_id", 0)

    for upd in updates:
        max_update_id = max(max_update_id, upd["update_id"])
        msg = upd.get("message")
        if not msg or "text" not in msg:
            continue

        chat_id = msg["chat"]["id"]
        text = msg["text"].strip()

        if text in ("/start", "/yordam", "/help"):
            send_message(
                chat_id,
                "Salom! Ertangi kun uchun rejani shu formatda (4 qatorda) yuboring:\n\n"
                "19:00\nAziz\nUy\nBozor\n\n"
                "Bir nechta kishi birga ketsa, Ism qatoriga vergul bilan yozing:\n"
                "19:00\nAziz, Vali\nUy\nBozor\n\n"
                "Xabar kechqurun soat 20:00 atrofida qayta ishlanadi, javob darhol kelmaydi.",
            )
            continue

        if text == "/owner":
            state["owner_chat_id"] = chat_id
            send_message(
                chat_id,
                "Siz asosiy qabul qiluvchi sifatida belgilandingiz. "
                "Kunlik Excel/Word fayllar shu chatga keladi.",
            )
            continue

        trip = parse_trip_message(text)
        if trip:
            trip["sender_chat_id"] = chat_id
            trips.append(trip)
        elif not text.startswith("/"):
            send_message(
                chat_id,
                "Xabar formati tushunarsiz. To'g'ri namuna uchun /yordam buyrug'ini yuboring.",
            )

    state["last_update_id"] = max_update_id
    return trips


def main():
    if not TOKEN:
        raise SystemExit("TELEGRAM_BOT_TOKEN muhit o'zgaruvchisi topilmadi.")

    state = load_state()
    updates = get_updates(state.get("last_update_id", 0))
    trips = process_updates(state, updates)

    owner_chat_id = state.get("owner_chat_id")

    if trips and owner_chat_id:
        merged = merge_trips(trips)
        assigned = assign_cars(merged, state["cars"], state.get("trip_duration_minutes", 45))

        tomorrow = (datetime.now(TASHKENT) + timedelta(days=1)).strftime("%Y-%m-%d")
        xlsx_path = f"kunlik_reja_{tomorrow}.xlsx"
        docx_path = f"kunlik_reja_{tomorrow}.docx"

        build_excel(assigned, state["cars"], tomorrow, xlsx_path)
        build_word(assigned, state["cars"], tomorrow, docx_path)

        for sender_chat_id in {t["sender_chat_id"] for t in trips}:
            send_message(sender_chat_id, f"Qabul qilindi. Ertangi ({tomorrow}) reja tayyorlandi.")

        send_document(owner_chat_id, xlsx_path, caption=f"{tomorrow} - kunlik reja (Excel)")
        send_document(owner_chat_id, docx_path, caption=f"{tomorrow} - kunlik reja (Word)")

        os.remove(xlsx_path)
        os.remove(docx_path)
    elif trips and not owner_chat_id:
        print("owner_chat_id aniqlanmagan - avval /owner buyrug'i yuborilishi kerak.")

    save_state(state)


if __name__ == "__main__":
    main()
